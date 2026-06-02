"""PDF-to-Word conversion module."""

from __future__ import annotations

import os
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path


@dataclass
class ConversionResult:
    """Result returned by :func:`convert_pdf_to_docx`."""

    input_path: Path
    output_path: Path
    pages: int


def _strip_watermarks(source: Path, destination: Path) -> None:
    """Write a copy of *source* to *destination* with background watermarks removed.

    This is a best-effort operation.  It removes path drawings that are both
    large (covering more than 40 % of the page area) **and** semi-transparent
    (fill or stroke opacity below 0.3), which is the most common pattern for
    background watermarks applied programmatically to PDFs.  Fully opaque
    large shapes are left untouched to avoid accidentally removing page
    backgrounds that are part of the real content.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(str(source))
    try:
        for page in doc:
            page_area = page.rect.width * page.rect.height
            if page_area <= 0:
                continue

            redact_rects: list[fitz.Rect] = []
            for drawing in page.get_drawings():
                rect = drawing.get("rect")
                if rect is None:
                    continue
                fill_opacity = drawing.get("fill_opacity")
                stroke_opacity = drawing.get("stroke_opacity")
                # Require at least one opacity value to be present and low
                opacities = [op for op in (fill_opacity, stroke_opacity) if op is not None]
                if not opacities:
                    continue
                if all(op >= 0.3 for op in opacities):
                    continue
                drawing_area = rect.width * rect.height
                if drawing_area > 0.4 * page_area:
                    redact_rects.append(fitz.Rect(rect))

            for r in redact_rects:
                page.add_redact_annot(r)
            if redact_rects:
                page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

        doc.save(str(destination), garbage=4, deflate=True)
    finally:
        doc.close()


def _do_convert(source: Path, destination: Path) -> None:
    """Invoke pdf2docx to convert *source* (PDF) to *destination* (DOCX)."""
    from pdf2docx import Converter  # type: ignore[import]

    cv = Converter(str(source))
    try:
        cv.convert(str(destination))
    finally:
        cv.close()


def default_word_output_path(input_path: str | Path) -> Path:
    """Return ``<stem>.docx`` next to the input PDF file."""
    return Path(input_path).with_suffix(".docx")


def is_scanned_pdf(path: str | Path, *, text_threshold: int = 10) -> bool:
    """Return ``True`` if the PDF appears to contain only scanned (image-based) pages.

    A page is considered scanned when it contains fewer than *text_threshold*
    characters of embedded text.  If every page in the document is below this
    threshold the PDF is classified as scanned.

    Args:
        path: Path to the PDF file to inspect.
        text_threshold: Minimum number of text characters required for a page
            to be considered as having a text layer.  Defaults to 10.

    Returns:
        ``True`` when all pages have minimal embedded text, ``False`` when at
        least one page contains a text layer.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    try:
        if doc.page_count == 0:
            return False
        for page in doc:
            if len(page.get_text().strip()) >= text_threshold:
                return False
        return True
    finally:
        doc.close()


def _ocr_pdf_to_docx(source: Path, destination: Path) -> None:
    """Convert *source* PDF to *destination* DOCX with OCR for scanned pages.

    Pages that already have an embedded text layer are converted using that
    text directly.  Pages with no embedded text (scanned pages) are rendered
    as high-resolution images and processed with EasyOCR.

    Args:
        source: Path to the input PDF.
        destination: Path for the output DOCX file.

    Raises:
        ImportError: If ``easyocr`` or ``python-docx`` are not installed.
    """
    try:
        import easyocr  # type: ignore[import]
    except ImportError as exc:
        raise ImportError(
            "OCR support requires the 'easyocr' package. "
            "Install it with:  pip install \"pdf-tools[ocr]\""
        ) from exc

    try:
        from docx import Document  # type: ignore[import]
    except ImportError as exc:
        raise ImportError(
            "OCR conversion requires the 'python-docx' package. "
            "Install it with:  pip install python-docx"
        ) from exc

    import numpy as np
    import fitz  # PyMuPDF

    reader = easyocr.Reader(["en"], verbose=False)

    doc_pdf = fitz.open(str(source))
    doc_word = Document()

    try:
        total = doc_pdf.page_count
        for page_num, page in enumerate(doc_pdf):
            text = page.get_text().strip()

            if text:
                # Page already has a text layer – use it directly.
                for line in text.splitlines():
                    stripped = line.strip()
                    if stripped:
                        doc_word.add_paragraph(stripped)
            else:
                # Scanned page – render to an image and run OCR.
                print(
                    f"  OCR processing page {page_num + 1}/{total}...",
                    file=sys.stderr,
                    flush=True,
                )
                pix = page.get_pixmap(dpi=300)
                # Build a NumPy array from the raw pixel samples to avoid a
                # PNG encode/decode round-trip.
                img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                    pix.height, pix.width, pix.n
                )
                results = reader.readtext(img_array)
                for _, ocr_text, _ in results:
                    stripped = ocr_text.strip()
                    if stripped:
                        doc_word.add_paragraph(stripped)

            if page_num < total - 1:
                doc_word.add_page_break()
    finally:
        doc_pdf.close()

    doc_word.save(str(destination))


def convert_pdf_to_docx(
    input_path: str | Path,
    output_path: str | Path | None = None,
    *,
    strip_watermarks: bool = False,
    force: bool = False,
    use_ocr: bool = False,
) -> ConversionResult:
    """Convert a PDF file to a Microsoft Word ``.docx`` file.

    Args:
        input_path: Path to the source PDF.
        output_path: Destination ``.docx`` path.  Defaults to
            ``<input-stem>.docx`` next to the input file.
        strip_watermarks: When ``True``, attempt to remove common background
            watermark elements (large semi-transparent shapes) from the PDF
            before conversion.
        force: Allow overwriting an existing output file.
        use_ocr: When ``True``, use OCR (via EasyOCR) to extract text from
            scanned pages.  Pages that already contain a text layer are
            converted using that text directly without OCR.  Requires the
            ``easyocr`` optional dependency (``pip install "pdf-tools[ocr]"``).

    Returns:
        A :class:`ConversionResult` with paths and page count.

    Raises:
        FileNotFoundError: If *input_path* does not exist or the output
            directory does not exist.
        ValueError: If *input_path* is not a file.
        FileExistsError: If *output_path* already exists and *force* is
            ``False``.
        ImportError: If *use_ocr* is ``True`` but ``easyocr`` is not
            installed.
    """
    import fitz  # PyMuPDF

    source = Path(input_path)
    if not source.exists():
        raise FileNotFoundError(f"Input file not found: {source}")
    if not source.is_file():
        raise ValueError(f"Input path is not a file: {source}")

    # Read the page count up front while the file is open for validation.
    doc = fitz.open(str(source))
    pages = doc.page_count
    doc.close()

    destination = Path(output_path) if output_path else default_word_output_path(source)

    if not destination.parent.exists():
        raise FileNotFoundError(
            f"Output directory does not exist: {destination.parent}"
        )
    if destination.exists() and not force:
        raise FileExistsError(
            f"Output file already exists: {destination}. "
            "Use --force / force=True to overwrite."
        )

    if strip_watermarks:
        tmp_fd, tmp_name = tempfile.mkstemp(suffix=".pdf")
        tmp_path = Path(tmp_name)
        try:
            os.close(tmp_fd)
            _strip_watermarks(source, tmp_path)
            if use_ocr:
                _ocr_pdf_to_docx(tmp_path, destination)
            else:
                _do_convert(tmp_path, destination)
        finally:
            tmp_path.unlink(missing_ok=True)
    elif use_ocr:
        _ocr_pdf_to_docx(source, destination)
    else:
        _do_convert(source, destination)

    return ConversionResult(
        input_path=source,
        output_path=destination,
        pages=pages,
    )
